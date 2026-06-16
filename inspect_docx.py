import docx

doc = docx.Document('CIMTRA-Legislativo.docx')
with open('docx_output.txt', 'w', encoding='utf-8') as f:
    f.write(f"Total tables: {len(doc.tables)}\n")
    for i, table in enumerate(doc.tables[:10]): 
        f.write(f"\n--- Table {i} ---\n")
        for j, row in enumerate(table.rows[:10]):
            texts = [cell.text.strip().replace('\n', ' ')[:50] for cell in row.cells]
            f.write(f"  Row {j}: {texts}\n")
